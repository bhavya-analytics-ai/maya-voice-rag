"""
ingest.py
Loads CSVs, PDFs, and Word docs from /context into ChromaDB.
Run once: python ingest.py
"""

import os
import pandas as pd
from dotenv import load_dotenv
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings
from langchain_chroma import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter

load_dotenv()

CONTEXT_DIR = "./context"
CHROMA_DIR = "./chroma_db"

splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)


# ── CSV loader ────────────────────────────────────────────────────────────────
def load_csv(filepath, filename):
    docs = []
    try:
        df = pd.read_csv(filepath)
    except Exception as e:
        print(f"  Error reading {filename}: {e}")
        return docs

    for _, row in df.iterrows():
        parts = []
        if "Product Name" in row and pd.notna(row["Product Name"]):
            parts.append(f"Product: {row['Product Name']}")
        if "SKU" in row and pd.notna(row["SKU"]):
            parts.append(f"SKU: {row['SKU']}")
        if "Stars" in row and pd.notna(row["Stars"]):
            parts.append(f"Rating: {row['Stars']} stars")
        if "MarketPlace" in row and pd.notna(row["MarketPlace"]):
            parts.append(f"Marketplace: {row['MarketPlace']}")
        if "Sentiment_GPT" in row and pd.notna(row["Sentiment_GPT"]):
            parts.append(f"Sentiment: {row['Sentiment_GPT']}")
        if "Review_English" in row and pd.notna(row["Review_English"]):
            parts.append(f"Review: {row['Review_English']}")
        if "Negative_Reason" in row and pd.notna(row["Negative_Reason"]):
            parts.append(f"Negative Reason: {row['Negative_Reason']}")
        if "Positive_Reason" in row and pd.notna(row["Positive_Reason"]):
            parts.append(f"Positive Reason: {row['Positive_Reason']}")
        if "Date" in row and pd.notna(row["Date"]):
            parts.append(f"Date: {row['Date']}")

        content = "\n".join(parts)
        if not content.strip():
            continue

        docs.append(Document(
            page_content=content,
            metadata={
                "source": filename,
                "type": "review",
                "product": str(row.get("Product Name", "Unknown")),
                "sku": str(row.get("SKU", "Unknown")),
                "sentiment": str(row.get("Sentiment_GPT", "Unknown")),
                "marketplace": str(row.get("MarketPlace", "Unknown")),
                "stars": str(row.get("Stars", "Unknown")),
            }
        ))
    return docs


# ── PDF loader ────────────────────────────────────────────────────────────────
def _ocr_pdf_with_vision(filepath, filename):
    """Convert image-based PDF pages to images and extract text via GPT-4o vision."""
    import fitz
    import base64
    from openai import OpenAI

    client = OpenAI()
    doc = fitz.open(filepath)
    all_text = ""

    for page_num, page in enumerate(doc):
        mat = fitz.Matrix(2, 2)
        pix = page.get_pixmap(matrix=mat)
        img_bytes = pix.tobytes("png")
        b64 = base64.b64encode(img_bytes).decode("utf-8")

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "This is a page from a Corona toilet product specification sheet. "
                            "Extract ALL text visible on this page exactly as it appears — "
                            "product name, SKU, dimensions, water consumption, weight, technical specs, "
                            "certifications, installation instructions, warnings. "
                            "Output plain text only, no markdown, preserve structure."
                        )
                    },
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/png;base64,{b64}"}
                    }
                ]
            }],
            max_tokens=2000,
        )
        page_text = response.choices[0].message.content.strip()
        all_text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"

    doc.close()
    return all_text


def load_pdf(filepath, filename):
    """Extract text from PDF — tries PyMuPDF first, falls back to GPT-4o vision for image-based PDFs."""
    try:
        import fitz
        doc = fitz.open(filepath)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()

        if len(text.strip()) < 500:
            print(f"  Image-based PDF detected — using GPT-4o vision OCR for {filename}...")
            text = _ocr_pdf_with_vision(filepath, filename)

        if not text.strip():
            print(f"  Warning: no text extracted from {filename}")
            return []

        chunks = splitter.split_text(text)
        print(f"  Extracted {len(chunks)} chunks from {filename}")
        return [Document(
            page_content=chunk,
            metadata={"source": filename, "type": "product_doc"}
        ) for chunk in chunks]
    except Exception as e:
        print(f"  Error reading PDF {filename}: {e}")
        return []


# ── Word doc loader ───────────────────────────────────────────────────────────
def load_docx(filepath, filename):
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(filepath)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

        if not text.strip():
            print(f"  Warning: no text extracted from {filename}")
            return []

        chunks = splitter.split_text(text)
        return [Document(
            page_content=chunk,
            metadata={"source": filename, "type": "product_doc"}
        ) for chunk in chunks]
    except Exception as e:
        print(f"  Error reading DOCX {filename}: {e}")
        return []


# ── Main loader ───────────────────────────────────────────────────────────────
def load_all_documents(context_dir):
    documents = []

    for filename in os.listdir(context_dir):
        filepath = os.path.join(context_dir, filename)
        ext = filename.lower().split(".")[-1]

        print(f"Loading {filename}...")

        if ext == "csv":
            documents.extend(load_csv(filepath, filename))
        elif ext == "pdf":
            documents.extend(load_pdf(filepath, filename))
        elif ext in ("docx", "doc"):
            documents.extend(load_docx(filepath, filename))
        else:
            print(f"  Skipping unsupported file: {filename}")

    return documents


def build_vectorstore(documents):
    print(f"\nEmbedding {len(documents)} documents into ChromaDB...")
    embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
    vectorstore = Chroma.from_documents(
        documents=documents,
        embedding=embeddings,
        persist_directory=CHROMA_DIR,
    )
    print(f"Done! Vector store saved to {CHROMA_DIR}")
    return vectorstore


if __name__ == "__main__":
    if not os.path.exists(CONTEXT_DIR):
        os.makedirs(CONTEXT_DIR)
        print(f"Created {CONTEXT_DIR}/ — add your files and run again.")
    else:
        docs = load_all_documents(CONTEXT_DIR)
        if not docs:
            print("No documents found.")
        else:
            build_vectorstore(docs)
            print(f"\nIngestion complete! {len(docs)} chunks indexed.")